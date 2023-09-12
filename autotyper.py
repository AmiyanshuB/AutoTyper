import tkinter as tk
import pyautogui
import time
import threading
from PIL import Image, ImageTk
from tkinter import filedialog, messagebox
import keyboard
from tkinter import messagebox
from PIL import Image, ImageTk


class AutoTyper:

    def __init__(self):
        self.is_running = True
        self.is_typing = False
        self.app = tk.Tk()
        self.app.title("Auto Typer")
        self.app.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.frame = tk.Frame(self.app)
        self.frame.pack(pady=20, padx=20)
        
        self.build_gui()


    def on_closing(self):
        self.is_running = False
        self.app.destroy()

    def start_typing(self):
        if not self.is_typing:
            self.is_typing = True
            thread = threading.Thread(target=self.auto_type)
            thread.start()

    def stop_typing(self):
        self.is_typing = False

    def auto_type(self):
        text = self.text_input.get("1.0", tk.END)
        start_delay = float(self.start_delay_spinbox.get())
        char_delay = float(self.char_delay_spinbox.get())
        word_delay = float(self.word_delay_spinbox.get())

        self.start_button.config(state=tk.DISABLED)
        print(f"Starting in {start_delay} seconds...")
        time.sleep(start_delay)

        word = ''
        for char in text:
            if not self.is_running or not self.is_typing or keyboard.is_pressed('esc') or keyboard.is_pressed('F10'):
                print("Stopped typing!")
                self.is_typing = False
                break
            
            if char == ' ':
                pyautogui.typewrite(word + ' ', interval=char_delay)
                word = ''
                if word_delay > 0:
                    time.sleep(word_delay)
            else:
                word += char

        if word:
            pyautogui.typewrite(word, interval=char_delay)

        self.start_button.config(state=tk.NORMAL)
        self.is_typing = False

    def build_gui(self):
        logo = Image.open("3824757.png")
        desired_size = (50, 50)  # Adjust as necessary
        logo = logo.resize(desired_size, resample=Image.LANCZOS)
        logo = ImageTk.PhotoImage(logo)
        self.logo_label = tk.Label(self.frame, image=logo, bg="gray")
        self.logo_label.image = logo

    # Place the logo to the left
        self.logo_label.grid(row=0, column=0, padx=(10, 5), pady=20)

    # Title (now placed to the right of the logo)
        self.label = tk.Label(self.frame, text="AutoTyper", font=("Arial", 24, "bold"), bg="gray")
        self.label.grid(row=0, column=1, pady=20, sticky="w")
        self.frame = tk.Frame(self.app, bg="gray")
        self.frame.pack(padx=10, pady=10, expand=True, fill="both")

    # Title
        self.label = tk.Label(self.frame, text="AutoTyper", font=("Arial", 24, "bold"), bg="gray")
        self.label.grid(row=0, column=0, columnspan=6, pady=(20, 10), sticky="w")

    # Text Input
        self.text_input = tk.Text(self.frame, height=10, wrap=tk.WORD)
        self.text_input.grid(row=1, column=0, columnspan=6, padx=10, pady=10, sticky="nsew")

    # Instructions for hotkeys
        self.instructions_label = tk.Label(self.frame, text="F9 to Start, F10 to Stop, Escape to Exit", font=("Arial", 12, "italic"), bg="gray", anchor="center")
        self.instructions_label.grid(row=2, column=0, columnspan=6, pady=10)

    # Configure the columns
        self.frame.grid_columnconfigure(0, weight=1)  # Left column for 'Import' button
        self.frame.grid_columnconfigure(1, weight=2)  # Middle column for 'Start' button
        self.frame.grid_columnconfigure(2, weight=1)  # Right column for 'Export' button

    # Import Button - Floating to the Left
        self.import_button = tk.Button(self.frame, text="Import", command=self.import_text, bg="green", fg="white")
        self.import_button.grid(row=3, column=0, padx=10, pady=5, sticky="w")

    # Start Button - Centered
        self.start_button = tk.Button(self.frame, text="Start Typing", command=self.start_typing, bg="blue", fg="white")
        self.start_button.grid(row=3, column=1, padx=10, pady=5)

    # Export Button - Floating to the Right
        self.export_button = tk.Button(self.frame, text="Export", command=self.export_text, bg="red", fg="white")
        self.export_button.grid(row=3, column=2, padx=10, pady=5, sticky="e")

    # Ensure the text input expands to fill the available vertical space
        self.frame.grid_rowconfigure(1, weight=1)



        keyboard.add_hotkey('F9', self.start_typing)
        keyboard.add_hotkey('F10', self.stop_typing)
        keyboard.add_hotkey('esc', self.stop_typing)
    def import_text(self):
        file_path = tk.filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx"), ("All Files", "*.*")])
        if file_path:
            if file_path.endswith('.txt'):
                with open(file_path, 'r') as file:
                    content = file.read()
                    self.text_input.delete("1.0", tk.END)  # Clear existing text
                    self.text_input.insert(tk.END, content)
            elif file_path.endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(file_path)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                    self.text_input.delete("1.0", tk.END)
                    self.text_input.insert(tk.END, content)
                except ImportError:
                    tk.messagebox.showerror("Error", "Please install python-docx to import Word documents.")
            else:
                tk.messagebox.showerror("Error", "Unsupported file type.")

    def export_text(self):
        file_path = tk.filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx"), ("All Files", "*.*")])
        if file_path:
            content = self.text_input.get("1.0", tk.END).strip()
            if file_path.endswith('.txt'):
                with open(file_path, 'w') as file:
                    file.write(content)
            elif file_path.endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(file_path)
                except ImportError:
                    tk.messagebox.showerror("Error", "Please install python-docx to export to Word documents.")
                else:
                    tk.messagebox.showerror("Error", "Unsupported file type.")



    def run(self):
        self.app.mainloop()


